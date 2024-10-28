from abc import ABC, abstractmethod
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
import os


class FileHandler(ABC):
    @abstractmethod
    def read_file(self, file_path):
        pass

    @abstractmethod
    def save_file(self, file_path, content):
        pass


class PDFHandler(FileHandler):
    def read_file(self, file_path):
        reader = PdfReader(file_path)
        pages = []
        for page in reader.pages:
            pages.append(page.extract_text())
        return pages

    def save_file(self, file_path, content):
        writer = PdfWriter()
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        for page_text in content:
            # Create a new PDF page with the translated text
            page = writer.add_blank_page(width=612, height=792)  # Standard letter size
            page.insert_text(12, 720, page_text)  # Basic text insertion
        with open(file_path, 'wb') as output_file:
            writer.write(output_file)


class DOCXHandler(FileHandler):
    def read_file(self, file_path):
        doc = Document(file_path)
        pages = []
        current_page = []
        for para in doc.paragraphs:
            if len(current_page) >= 3000:  # Arbitrary page size
                pages.append('\n'.join(current_page))
                current_page = []
            current_page.append(para.text)
        if current_page:
            pages.append('\n'.join(current_page))
        return pages

    def save_file(self, file_path, content):
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        doc = Document()
        for page_text in content:
            paragraphs = page_text.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    doc.add_paragraph(para_text)
        doc.save(file_path)


class TXTHandler(FileHandler):
    def read_file(self, file_path):
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        # Split into pages based on some criteria (e.g., line count)
        return [content]  # For now, treating as single page

    def save_file(self, file_path, content):
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(''.join(content))
